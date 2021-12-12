from os.path import join
from time import time_ns

from docx import Document
from docx.shared import Inches
from data import Campain, Screen, Combat
from data.base import projectRoot
from rich.console import Console


def export(campaign:int):

    campaignData = Campain()
    screenData = Screen()
    con = Console()

    docObj = Document()
    
    if str(campaign) not in campaignData.readDoc_ids():
        raise TypeError('doc id not found')

    # get data
    cData = campaignData.readById(campaign)
    sData = screenData.readBycampain(campaign)

    docObj.add_heading(cData['title'], 0)
    docObj.add_paragraph(cData['bio'])
    docObj.add_page_break()
    docObj.add_heading("Screens", 1)

    con.print('exporting screens')

    for screen in sData:

        docObj.add_heading(
            screen['title'],
            level=2
        )

        docObj.add_picture(
            screen['picture'],
            width=Inches(5.0)
        )

        # docObj.add_page_break()

        tbl = docObj.add_table(rows=2, cols=2)
        firstCells = tbl.rows[0].cells
        firstCells[0].text = "Player Notes"
        firstCells[1].text = "DM's Notes"

        secoundCells = tbl.rows[1].cells
        secoundCells[0].text = screen['pl_notes']
        secoundCells[1].text = screen['dm_notes']   

        docObj.add_page_break()

    docObj.add_heading("Encounters", 1)

    tbl = docObj.add_table(rows=1, cols=3)
    firstCells = tbl.rows[0].cells
    firstCells[0].text = "Encounter ID"
    firstCells[1].text = "name"
    firstCells[2].text = "url"

    con.print('exporting Encounters')

    for row in Combat().readByCompagn(campaign):
        rows = tbl.add_row().cells
        rows[0].text = str(row.doc_id)
        rows[1].text = row['name']
        rows[2].text = row['url']
    ts = time_ns()
    path = join(projectRoot(), '.local/', f'{campaign}-export-{ts}.docx')
    con.print(f'exported to {path}')
    docObj.save(path)
